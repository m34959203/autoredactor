#!/usr/bin/env python3
"""
Авторедактор журнала
Автоматизирует сборку журнала из статей авторов в Word формате
"""

import os
import re
import sys
from pathlib import Path
from typing import Optional
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from pypdf import PdfReader, PdfWriter
import unicodedata


class JournalEditor:
    """Класс для автоматической сборки журнала"""
    
    def __init__(self, 
                 articles_dir: str,
                 output_dir: str,
                 title_page: Optional[str] = None,
                 first_pages: Optional[str] = None,
                 end_pages: Optional[str] = None):
        """
        Инициализация редактора
        
        Args:
            articles_dir: Папка со статьями в формате .docx
            output_dir: Папка для выходных файлов
            title_page: PDF титульного листа (опционально)
            first_pages: PDF с 1-2 страницами (опционально)
            end_pages: PDF с последними страницами (опционально)
        """
        self.articles_dir = Path(articles_dir)
        self.output_dir = Path(output_dir)
        self.title_page = Path(title_page) if title_page else None
        self.first_pages = Path(first_pages) if first_pages else None
        self.end_pages = Path(end_pages) if end_pages else None
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Регистрация шрифтов с поддержкой кириллицы
        self._register_fonts()
        
        # Стили для PDF
        self.styles = self._create_styles()
        
        # Данные для содержания
        self.toc_entries = []  # [(название, начальная_страница), ...]
        
    def _register_fonts(self):
        """Регистрация шрифтов с поддержкой кириллицы"""
        font_paths = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
            '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                font_name = Path(path).stem
                try:
                    pdfmetrics.registerFont(TTFont(font_name, path))
                except:
                    pass
        
        # Проверяем, какие шрифты доступны
        self.font_name = 'DejaVuSans'
        self.font_bold = 'DejaVuSans-Bold'
        
    def _create_styles(self):
        """Создание стилей для PDF"""
        styles = getSampleStyleSheet()
        
        # Стиль для основного текста статьи
        styles.add(ParagraphStyle(
            name='ArticleBody',
            fontName=self.font_name,
            fontSize=11,
            leading=14,
            firstLineIndent=1*cm,  # Отступ первой строки
            spaceBefore=6,
            spaceAfter=6,
        ))
        
        # Стиль для заголовка статьи
        styles.add(ParagraphStyle(
            name='ArticleTitle',
            fontName=self.font_bold,
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=12,
            alignment=1,  # Центр
        ))
        
        # Стиль для содержания - заголовок
        styles.add(ParagraphStyle(
            name='TOCTitle',
            fontName=self.font_bold,
            fontSize=16,
            leading=20,
            spaceBefore=20,
            spaceAfter=20,
            alignment=1,
        ))
        
        # Стиль для элемента содержания
        styles.add(ParagraphStyle(
            name='TOCEntry',
            fontName=self.font_name,
            fontSize=11,
            leading=14,
            spaceBefore=4,
            spaceAfter=4,
        ))
        
        return styles
    
    def _is_latin(self, char: str) -> bool:
        """Проверка, является ли символ латинским"""
        try:
            return 'LATIN' in unicodedata.name(char)
        except:
            return False
    
    def _is_cyrillic(self, char: str) -> bool:
        """Проверка, является ли символ кириллическим"""
        try:
            return 'CYRILLIC' in unicodedata.name(char)
        except:
            return False
    
    def _get_sort_key(self, title: str) -> tuple:
        """
        Ключ сортировки: латиница в начале, потом кириллица
        Возвращает кортеж (группа, нормализованное_название)
        """
        title_clean = title.strip().upper()
        first_char = title_clean[0] if title_clean else ''
        
        if self._is_latin(first_char):
            group = 0  # Латиница первая
        elif self._is_cyrillic(first_char):
            group = 1  # Кириллица вторая
        else:
            group = 2  # Остальное в конце
            
        return (group, title_clean)
    
    def _extract_title_from_docx(self, docx_path: Path) -> str:
        """Извлечение названия статьи из документа Word"""
        try:
            doc = DocxDocument(str(docx_path))
            # Берём первый непустой параграф как название
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Ограничиваем длину названия
                    if len(text) > 200:
                        text = text[:200] + "..."
                    return text
            return docx_path.stem  # Если нет текста, используем имя файла
        except Exception as e:
            print(f"Ошибка чтения {docx_path}: {e}")
            return docx_path.stem
    
    def _extract_text_from_docx(self, docx_path: Path) -> list:
        """Извлечение текста из документа Word"""
        paragraphs = []
        try:
            doc = DocxDocument(str(docx_path))
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
        except Exception as e:
            print(f"Ошибка чтения {docx_path}: {e}")
        return paragraphs
    
    def get_sorted_articles(self) -> list:
        """
        Получение списка статей, отсортированных по алфавиту
        Латиница в начале, затем кириллица
        """
        articles = []
        
        # Находим все .docx файлы
        docx_files = list(self.articles_dir.glob('*.docx'))
        
        for docx_path in docx_files:
            title = self._extract_title_from_docx(docx_path)
            articles.append({
                'path': docx_path,
                'title': title,
                'sort_key': self._get_sort_key(title)
            })
        
        # Сортировка: латиница, затем кириллица, по алфавиту внутри группы
        articles.sort(key=lambda x: x['sort_key'])
        
        return articles
    
    def _convert_article_to_pdf(self, article: dict, output_path: Path) -> int:
        """
        Конвертация статьи в PDF с 4 отступами (абзацами)
        Возвращает количество страниц
        """
        paragraphs = self._extract_text_from_docx(article['path'])
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        story = []
        
        # 4 отступа перед статьёй (пустые строки)
        for _ in range(4):
            story.append(Spacer(1, 12))
        
        # Заголовок статьи
        story.append(Paragraph(article['title'], self.styles['ArticleTitle']))
        story.append(Spacer(1, 12))
        
        # Текст статьи (пропускаем первый параграф - он уже использован как заголовок)
        for i, para_text in enumerate(paragraphs):
            if i == 0:
                continue  # Пропускаем заголовок
            # Экранируем HTML-спецсимволы
            safe_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_text, self.styles['ArticleBody']))
        
        # Добавляем разрыв страницы после статьи
        story.append(PageBreak())
        
        doc.build(story)
        
        # Считаем страницы
        reader = PdfReader(str(output_path))
        return len(reader.pages)
    
    def _create_toc_pdf(self, output_path: Path):
        """Создание PDF с содержанием"""
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        story = []
        
        # Заголовок содержания
        story.append(Paragraph("СОДЕРЖАНИЕ", self.styles['TOCTitle']))
        story.append(Spacer(1, 20))
        
        # Формируем таблицу содержания
        toc_data = []
        for title, page_num in self.toc_entries:
            # Ограничиваем длину названия для таблицы
            display_title = title if len(title) <= 70 else title[:67] + "..."
            toc_data.append([display_title, str(page_num)])
        
        if toc_data:
            table = Table(toc_data, colWidths=[14*cm, 2*cm])
            table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), self.font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(table)
        
        doc.build(story)
    
    def build_journal(self, output_filename: str = "journal.pdf"):
        """
        Сборка журнала
        
        Args:
            output_filename: Имя выходного файла
        """
        print("=" * 50)
        print("АВТОРЕДАКТОР ЖУРНАЛА")
        print("=" * 50)
        
        # 1. Получаем отсортированные статьи
        print("\n[1] Сортировка статей по алфавиту...")
        articles = self.get_sorted_articles()
        
        if not articles:
            print("ОШИБКА: Статьи не найдены в папке", self.articles_dir)
            return
        
        print(f"    Найдено статей: {len(articles)}")
        for i, art in enumerate(articles, 1):
            prefix = "LAT" if art['sort_key'][0] == 0 else "CYR"
            print(f"    {i}. [{prefix}] {art['title'][:50]}...")
        
        # 2. Конвертируем статьи в PDF
        print("\n[2] Конвертация статей в PDF...")
        temp_pdfs = []
        current_page = 1
        
        # Учитываем титульный лист и первые страницы
        if self.title_page and self.title_page.exists():
            reader = PdfReader(str(self.title_page))
            current_page += len(reader.pages)
            
        if self.first_pages and self.first_pages.exists():
            reader = PdfReader(str(self.first_pages))
            current_page += len(reader.pages)
        
        for i, article in enumerate(articles):
            temp_pdf = self.output_dir / f"_temp_article_{i}.pdf"
            print(f"    Обработка: {article['title'][:40]}...")
            
            # Записываем номер страницы для содержания
            self.toc_entries.append((article['title'], current_page))
            
            # Конвертируем статью
            pages = self._convert_article_to_pdf(article, temp_pdf)
            temp_pdfs.append(temp_pdf)
            current_page += pages
        
        # 3. Создаём PDF с содержанием
        print("\n[3] Формирование содержания...")
        toc_pdf = self.output_dir / "_temp_toc.pdf"
        self._create_toc_pdf(toc_pdf)
        
        # 4. Собираем финальный PDF
        print("\n[4] Сборка журнала...")
        writer = PdfWriter()
        
        # Титульный лист
        if self.title_page and self.title_page.exists():
            print("    + Титульный лист")
            reader = PdfReader(str(self.title_page))
            for page in reader.pages:
                writer.add_page(page)
        
        # Первые страницы (1-2)
        if self.first_pages and self.first_pages.exists():
            print("    + Первые страницы")
            reader = PdfReader(str(self.first_pages))
            for page in reader.pages:
                writer.add_page(page)
        
        # Статьи
        print(f"    + Статьи ({len(temp_pdfs)} шт.)")
        for temp_pdf in temp_pdfs:
            reader = PdfReader(str(temp_pdf))
            for page in reader.pages:
                writer.add_page(page)
        
        # Содержание (после статей)
        print("    + Содержание")
        reader = PdfReader(str(toc_pdf))
        for page in reader.pages:
            writer.add_page(page)
        
        # Последние страницы
        if self.end_pages and self.end_pages.exists():
            print("    + Последние страницы")
            reader = PdfReader(str(self.end_pages))
            for page in reader.pages:
                writer.add_page(page)
        
        # Сохраняем
        output_path = self.output_dir / output_filename
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        # Очистка временных файлов
        print("\n[5] Очистка временных файлов...")
        for temp_pdf in temp_pdfs:
            temp_pdf.unlink(missing_ok=True)
        toc_pdf.unlink(missing_ok=True)
        
        print("\n" + "=" * 50)
        print(f"ГОТОВО! Журнал сохранён: {output_path}")
        print(f"Всего страниц: {len(writer.pages)}")
        print("=" * 50)
        
        return output_path


def main():
    """Точка входа"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Авторедактор журнала - автоматическая сборка журнала из статей',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  %(prog)s ./articles ./output
  %(prog)s ./articles ./output --title title.pdf --first pages_1_2.pdf --end last.pdf
  %(prog)s ./articles ./output -o my_journal.pdf
        """
    )
    
    parser.add_argument('articles_dir', help='Папка со статьями в формате .docx')
    parser.add_argument('output_dir', help='Папка для выходных файлов')
    parser.add_argument('-o', '--output', default='journal.pdf', help='Имя выходного файла (по умолчанию: journal.pdf)')
    parser.add_argument('--title', help='PDF титульного листа')
    parser.add_argument('--first', help='PDF с 1-2 страницами журнала')
    parser.add_argument('--end', help='PDF с последними страницами журнала')
    
    args = parser.parse_args()
    
    editor = JournalEditor(
        articles_dir=args.articles_dir,
        output_dir=args.output_dir,
        title_page=args.title,
        first_pages=args.first,
        end_pages=args.end
    )
    
    editor.build_journal(args.output)


if __name__ == '__main__':
    main()
