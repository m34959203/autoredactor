from docx import Document
from typing import Optional
import os


class DocxParser:
    """Service for parsing DOCX files."""

    @staticmethod
    def extract_text(file_path: str, max_chars: int = 2000) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file
            max_chars: Maximum number of characters to extract

        Returns:
            Extracted text from the beginning of the document
        """
        try:
            doc = Document(file_path)
            text_parts = []
            total_chars = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    total_chars += len(text)

                    if total_chars >= max_chars:
                        break

            return "\n".join(text_parts)[:max_chars]
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")

    @staticmethod
    def get_full_text(file_path: str) -> str:
        """
        Extract all text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Full text content
        """
        try:
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")

    @staticmethod
    def validate_docx(file_path: str) -> bool:
        """
        Validate that file is a valid DOCX document.

        Args:
            file_path: Path to file

        Returns:
            True if valid DOCX, False otherwise
        """
        try:
            Document(file_path)
            return True
        except Exception:
            return False
